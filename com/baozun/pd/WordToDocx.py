from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 你的单词列表（这里用示例数据，实际替换为你的100个单词）
words = [
    "brave /breɪv/", "negligible /'neɡlɪdʒəbl/", "logistics /lə'dʒɪstɪks/",
    "willpower /'wɪlpaʊə(r)/", "wag /wæɡ/", "villa /'vɪlə/",
    "vice /vaɪs/", "vaccinate /'væksɪneɪt/", "up-to-date /ˌʌp tə ˈdeɪt/",
    "tortoise /'tɔːtəs/", "temperate /'tempərət/", "swan /swɒn/",
    # 此处省略剩余单词，实际使用时替换为你的完整列表
    "nobility /nəʊ'bɪləti/", "decent /'diːsnt/", "exceed /ɪk'siːd/"
]

# 确保取前100个单词（防止列表过长）
words = words[:100]


def create_two_column_word(words_list, output_file="单词两列排版.docx"):
    """
    将单词列表按两列对齐写入docx文件
    :param words_list: 单词列表
    :param output_file: 输出文件名
    """
    # 创建新的docx文档
    doc = Document()

    # 设置文档标题
    title = doc.add_heading('单词两列排版', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 标题居中

    # 计算每列的单词数量
    total_words = len(words_list)
    col1_count = (total_words + 1) // 2  # 第一列数量（奇数时多一个）
    col2_count = total_words // 2  # 第二列数量

    # 分割单词列表为两列
    col1_words = words_list[:col1_count]
    col2_words = words_list[col1_count:]

    # 创建表格实现两列对齐（推荐方式，对齐效果最好）
    table = doc.add_table(rows=col1_count, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体居中
    table.style = 'Table Grid'  # 显示表格边框（可选，便于查看）

    # 设置表格列宽
    for row in table.rows:
        row.cells[0].width = Inches(3)  # 第一列宽度
        row.cells[1].width = Inches(3)  # 第二列宽度

    # 填充表格内容
    for i in range(col1_count):
        # 获取当前行的两个单元格
        cell1 = table.cell(i, 0)
        cell2 = table.cell(i, 1)

        # 设置单元格文字格式
        para1 = cell1.paragraphs[0]
        run1 = para1.add_run(col1_words[i])
        run1.font.size = Pt(8)  # 字体大小
        run1.font.name = '宋体'  # 字体

        # 第二列如果有单词则填充
        if i < col2_count:
            para2 = cell2.paragraphs[0]
            run2 = para2.add_run(col2_words[i])
            run2.font.size = Pt(8)
            run2.font.name = '宋体'

        # 设置单元格文字居中对齐
        para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if i < col2_count:
            para2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 保存文档
    doc.save(output_file)
    print(f"✅ 文档已生成：{output_file}")
    print(f"📊 排版统计：总单词数={total_words}，第一列={col1_count}，第二列={col2_count}")


if __name__ == "__main__":
    # 执行排版
    with open(r'D:\u2026\所有生词本_20260222_1903.txt', 'r', encoding='utf-8') as file:
        txt = file.readlines()
        rs = []
        for item in txt:
            rs.append(item.replace('\n', ''))
        create_two_column_word(rs, output_file=r'D:\u02\json2Excel\test.docx')